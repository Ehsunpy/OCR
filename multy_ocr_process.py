
import io
import os
import sys

import fitz  # PyMuPDF
import cv2
import numpy as np
from PIL import Image
import pytesseract
from concurrent.futures import ProcessPoolExecutor, as_completed
from multiprocessing import cpu_count
import os
from docx import Document
from tqdm import tqdm

# مسیر تسرکت ویندوز
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
Image.MAX_IMAGE_PIXELS = None
def ocr_from_image(
    image_path: str,
    lang: str = "fas+eng",
        tess_config="--oem 1 --psm 6 -c preserve_interword_spaces=1"
):
    from PIL import Image
    import cv2, numpy as np

    pil_img = Image.open(image_path)
    # پیش‌پردازش مختصر
    img_cv = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (3,3), 0)
    _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY+cv2.THRESH_OTSU)
    proc = Image.fromarray(thresh)

    return pytesseract.image_to_string(proc, lang=lang, config=tess_config)

def preprocess_pil(pil_img: Image.Image) -> Image.Image:
    """پیش‌پردازش: خاکستری، بلور ملایم، اوتسو باینری"""
    Image.MAX_IMAGE_PIXELS = None  # یعنی محدودیت رو بردار

    img_cv = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (3, 3), 0)
    _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    pil_out = Image.fromarray(thresh)

    # 🔹 کوچک‌سازی تصویر (مثلاً نصف عرض و ارتفاع)
    max_width = 2000  # هرچی خواستی بزار
    if pil_out.width > max_width:
        ratio = max_width / pil_out.width
        new_size = (int(pil_out.width * ratio), int(pil_out.height * ratio))
        pil_out = pil_out.resize(new_size, Image.LANCZOS)

    return pil_out


def ocr_worker(task):
    """
    ورودی: (page_idx, png_bytes, lang, tess_config)
    خروجی: (page_idx, text)
    """
    page_idx, png_bytes, lang, tess_config = task
    pil_img = Image.open(io.BytesIO(png_bytes))
    proc = preprocess_pil(pil_img)
    text = pytesseract.image_to_string(proc, lang=lang, )
    return page_idx, text

def page_to_png_bytes(page, dpi=300):
    """رندر یک صفحه به PNG (bytes) برای انتقال سبک به پروسس‌ها"""
    pix = page.get_pixmap(dpi=dpi)
    pil = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    buf = io.BytesIO()
    # PNG بدون افت کیفیت؛ برای OCR بهتر از JPEG است
    pil.save(buf, format="PNG")
    return buf.getvalue()

def ocr_from_pdf_multiproc(
    pdf_path: str,
    lang: str = "fas+eng",
    dpi: int = 200,
    workers: int | None = None,
    tess_config: str = "--oem 1 --psm 6"
):
    """
    رندر سری‌وار صفحات PDF → ارسال پیش‌پردازش+OCR به چند پروسس.
    - workers: پیش‌فرض min(8, cpu_count()-1)
    - tess_config: تنظیمات سرعت/چیدمان تسرکت
    """
    if workers is None:
        # معمولاً این انتخاب، سرعت خوب + عدم قفل سیستم
        workers = max(1, min(8, cpu_count() - 1))

    results = []
    futures = []

    with fitz.open(pdf_path) as doc, ProcessPoolExecutor(max_workers=workers) as ex:
        for page_idx in range(doc.page_count):
            pg = doc.load_page(page_idx)
            png_bytes = page_to_png_bytes(pg, dpi=dpi)
            fut = ex.submit(ocr_worker, (page_idx, png_bytes, lang, tess_config))
            futures.append(fut)

        for fut in as_completed(futures):
            page_idx, text = fut.result()
            results.append((page_idx, text))

    # مرتب‌سازی بر اساس شماره صفحه
    results.sort(key=lambda x: x[0])

    # چسباندن متن‌ها با هدر صفحه
    full_text = []
    for idx, txt in results:
        full_text.append(f"--- صفحه {idx+1} ---\n{txt}\n")
    return "".join(full_text)
if __name__ == "__main__":
    # from multy_ocr_process import ocr_from_pdf_multiproc

    base_dir =sys.argv[1]


    files = os.listdir(base_dir)
    paths = [f for f in files if f.lower().endswith(".pdf")]

    ocr_dir = os.path.join(base_dir, "OCR")
    os.makedirs(ocr_dir, exist_ok=True)

    for pdf in tqdm(paths, desc="Processing PDFs"):
        pdf_path = os.path.join(base_dir, pdf)
        all_text = ocr_from_pdf_multiproc(pdf_path, dpi=150)

        docx_filename = os.path.splitext(pdf)[0] + ".docx"
        docx_path = os.path.join(ocr_dir, docx_filename)

        document = Document()
        document.add_paragraph(all_text)
        document.save(docx_path)

        print(f"✅ OCR ذخیره شد: {docx_path}")


    #
    # for pdf in paths:
    #     pdf_path = os.path.join(base_dir, pdf)
    #     all_text = ocr_from_pdf_multiproc(pdf_path)
    #
    #     # اسم فایل Markdown
    #     md_filename = os.path.splitext(pdf)[0] + ".md"
    #     md_path = os.path.join(ocr_dir, md_filename)
    #
    #     # ذخیره به صورت Markdown
    #     with open(md_path, "w", encoding="utf-8") as f:
    #         for idx, page_text in enumerate(all_text.split("--- صفحه ")):
    #             if page_text.strip() == "":
    #                 continue
    #             # جداکننده صفحه و متن
    #             f.write(f"--- صفحه {page_text.strip()}\n\n")
    #
    #     print(f"✅ OCR ذخیره شد: {md_path}")
#
#
# import io
# import os
# import sys
#
# import fitz  # PyMuPDF
# import cv2
# import numpy as np
# from PIL import Image
# import pytesseract
# from concurrent.futures import ProcessPoolExecutor, as_completed
# from multiprocessing import cpu_count
# import os
# from docx import Document
# # مسیر تسرکت ویندوز
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# Image.MAX_IMAGE_PIXELS = None
# def ocr_from_image(
#     image_path: str,
#     lang: str = "fas+eng",
#     tess_config: str = "--oem 1 --psm 6"
# ):
#     from PIL import Image
#     import cv2, numpy as np
#
#     pil_img = Image.open(image_path)
#     # پیش‌پردازش مختصر
#     img_cv = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
#     gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
#     blurred = cv2.GaussianBlur(gray, (3,3), 0)
#     _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY+cv2.THRESH_OTSU)
#     proc = Image.fromarray(thresh)
#
#     return pytesseract.image_to_string(proc, lang=lang, config=tess_config)
#
# def preprocess_pil(pil_img: Image.Image) -> Image.Image:
#     """پیش‌پردازش: خاکستری، بلور ملایم، اوتسو باینری"""
#     Image.MAX_IMAGE_PIXELS = None  # یعنی محدودیت رو بردار
#
#     img_cv = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
#     gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
#     blurred = cv2.GaussianBlur(gray, (3, 3), 0)
#     _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
#
#     pil_out = Image.fromarray(thresh)
#
#     # 🔹 کوچک‌سازی تصویر (مثلاً نصف عرض و ارتفاع)
#     max_width = 2000  # هرچی خواستی بزار
#     if pil_out.width > max_width:
#         ratio = max_width / pil_out.width
#         new_size = (int(pil_out.width * ratio), int(pil_out.height * ratio))
#         pil_out = pil_out.resize(new_size, Image.LANCZOS)
#
#     return pil_out
#
#
# def ocr_worker(task):
#     """
#     ورودی: (page_idx, png_bytes, lang, tess_config)
#     خروجی: (page_idx, text)
#     """
#     page_idx, png_bytes, lang, tess_config = task
#     pil_img = Image.open(io.BytesIO(png_bytes))
#     proc = preprocess_pil(pil_img)
#     text = pytesseract.image_to_string(proc, lang=lang, )
#     return page_idx, text
#
# def page_to_png_bytes(page, dpi=300):
#     """رندر یک صفحه به PNG (bytes) برای انتقال سبک به پروسس‌ها"""
#     pix = page.get_pixmap(dpi=dpi)
#     pil = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
#     buf = io.BytesIO()
#     # PNG بدون افت کیفیت؛ برای OCR بهتر از JPEG است
#     pil.save(buf, format="PNG")
#     return buf.getvalue()
#
# def ocr_from_pdf_multiproc(
#     pdf_path: str,
#     lang: str = "fas+eng",
#     dpi: int = 200,
#     workers: int | None = None,
#     tess_config: str = "--oem 1 --psm 6"
# ):
#     """
#     رندر سری‌وار صفحات PDF → ارسال پیش‌پردازش+OCR به چند پروسس.
#     - workers: پیش‌فرض min(8, cpu_count()-1)
#     - tess_config: تنظیمات سرعت/چیدمان تسرکت
#     """
#     if workers is None:
#         # معمولاً این انتخاب، سرعت خوب + عدم قفل سیستم
#         workers = max(1, min(8, cpu_count() - 1))
#
#     results = []
#     futures = []
#
#     with fitz.open(pdf_path) as doc, ProcessPoolExecutor(max_workers=workers) as ex:
#         for page_idx in range(doc.page_count):
#             pg = doc.load_page(page_idx)
#             png_bytes = page_to_png_bytes(pg, dpi=dpi)
#             fut = ex.submit(ocr_worker, (page_idx, png_bytes, lang, tess_config))
#             futures.append(fut)
#
#         for fut in as_completed(futures):
#             page_idx, text = fut.result()
#             results.append((page_idx, text))
#
#     # مرتب‌سازی بر اساس شماره صفحه
#     results.sort(key=lambda x: x[0])
#
#     # چسباندن متن‌ها با هدر صفحه
#     full_text = []
#     for idx, txt in results:
#         full_text.append(f"--- صفحه {idx+1} ---\n{txt}\n")
#     return "".join(full_text)
# if __name__ == "__main__":
#     # from multy_ocr_process import ocr_from_pdf_multiproc
#
#     base_dir =sys.argv[1]
#
#     if os.path.isfile(base_dir):
#         ext = os.path.splitext(base_dir)[1].lower()
#         if ext == ".pdf":
#             text = ocr_from_pdf_multiproc(base_dir, lang="fas+eng", dpi=300, tess_config="--oem 1 --psm 6")
#             print(text)
#         else:
#             # تصویر (TIF/PNG/JPG…)
#             text = ocr_from_image(base_dir, lang="fas+eng", tess_config="--oem 1 --psm 6")
#             print(text)
#     files = os.listdir(base_dir)
#     paths = [f for f in files if f.lower().endswith(".pdf")]
#
#     ocr_dir = os.path.join(base_dir, "OCR")
#     os.makedirs(ocr_dir, exist_ok=True)
#
#     for pdf in paths:
#         pdf_path = os.path.join(base_dir, pdf)
#         all_text = ocr_from_pdf_multiproc(pdf_path)
#
#         docx_filename = os.path.splitext(pdf)[0] + ".docx"
#         docx_path = os.path.join(ocr_dir, docx_filename)
#
#         document = Document()
#         document.add_paragraph(all_text)
#         document.save(docx_path)
#
#         print(f"✅ OCR ذخیره شد: {docx_path}")
#
